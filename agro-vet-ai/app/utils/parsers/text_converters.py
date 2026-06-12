import logging
import os
import pandas as pd
import subprocess
from docx import Document
from app.utils.parsers.base_converter import BaseConverter
from typing import Optional


class TxtConverter(BaseConverter):
    """Конвертер текстовых файлов: возвращает содержимое фалйа .txt"""

    def convert(self) -> Optional[str]:
        """
        Читает и возвращает содержимое текстового файла.

        :return: Текст файла или None в случае ошибки.
        """
        try:
            with open(self.file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logging.exception(f"Ошибка при чтении файла {self.file_path}: {e}")
            return None


class DocxConverter(BaseConverter):
    """Конвертер Word-файлов (.doc, .docx). Конвертирует текст и таблицы в CSV."""

    def is_word_file(self):
        """
        Проверяет, является ли файл Word-документом (.doc или .docx).

        :return: True, если файл Word, иначе False.
        """
        return self.file_path.endswith('.doc') or self.file_path.endswith('.docx')

    def is_file_path(self):
        """
        Проверяет, имеет ли файл расширение .doc.

        :return: True, если файл .doc, иначе False.
        """
        return self.file_path.endswith('.doc')

    def convert_doc_to_docx(self) -> str:
        """
        Конвертирует файл .doc в .docx с использованием LibreOffice.

        :return: Путь к новому .docx файлу.
        """
        # Определить путь к выходному файлу
        output_file = self.file_path + "x"
        output_dir = os.path.dirname(self.file_path)

        # Запустить libreoffice с указанным выходным каталогом
        subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'docx', self.file_path, '--outdir', output_dir],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )

        # Возвращает путь к новому файлу .docx
        return output_file

    def convert(self) -> Optional[str]:
        """
        Конвертирует Word-файл в CSV-формат (текст и таблицы).

        :return: Данные в формате CSV или None в случае ошибки.
        """
        if not self.is_word_file():
            logging.exception(f"Файл {self.file_path} не является файлом Word (.doc или .docx).")
            return None

        try:
            if self.is_file_path():
                logging.info(f"Преобразование файла {self.file_path} из формата .doc в формат .docx")
                self.file_path = self.convert_doc_to_docx()
        except Exception as e:
            logging.exception(f"Ошибка при преобразовании файла {self.file_path} из .doc в .docx: {e}")
            return None

        document = Document(self.file_path)
        csv_data = []

        # Выполнить итерацию по всем таблицам в документе
        if not document.tables:
            for para in document.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    csv_data.append(para_text)

            df = pd.DataFrame(csv_data, columns=["Text"])
            csv_string = df.to_csv(index=False, sep=";")
            if csv_string is not None:
                csv_string = csv_string.strip()

            return csv_string
        else:
            for table in document.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    # Добавлять только непустые строки
                    if any(row_data):
                        table_data.append(row_data)

                if table_data:
                    df = pd.DataFrame(table_data)
                    df = df.dropna(how='all').loc[:, (df != "").any(axis=0)]
                    csv_string = df.to_csv(index=False, sep=";")
                    if csv_string is not None:
                        csv_string = csv_string.strip()

                    csv_data.append(csv_string)

        return "\n\n".join(csv_data)
